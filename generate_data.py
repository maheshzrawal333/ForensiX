"""
==============================================================================
Enterprise Forensic Data Synthesizer (E2E Testing Utility)
==============================================================================
ARCHITECTURAL OVERVIEW:
This utility dynamically generates a localized, highly structured, multi-format
investigation case. Its primary purpose is to test the Retrieval-Augmented
Generation (RAG) backend (specifically the Apache Tika parsers, Token Splitters,
and pgvector Hybrid Search logic).

TESTING STRATEGY (The Needle in the Haystack):
To prove the AI isn't hallucinating, we bury specific "Golden Clues" deep within
large files (e.g., Row 250 of a 500-line CSV, or Line 800 of a 1000-line TXT file).
The LLM must successfully extract, connect, and cite these exact clues to pass the
end-to-end integration test.
"""

import os
import random
import csv
from faker import Faker
from docx import Document
from reportlab.pdfgen import canvas
from datetime import datetime, timedelta

# Initialize Faker with a static seed to ensure deterministic generation across
# different developer machines. This guarantees the "noise" around the golden clues
# remains exactly the same on every run, preventing flaky E2E test failures.
fake = Faker()
Faker.seed(42)

# ==============================================================================
# INVESTIGATION TAXONOMY (The Ground Truth)
# ==============================================================================
# These variables represent the mathematical "Ground Truth" of our test case.
# The UI will be tasked with identifying these exact strings via semantic search.
CASE_NAME = "Operation_Midnight"
TARGET_ENTITIES = {
    "mastermind": "Victor Sterling",
    "logistics": "Elena Rostova",
    "muscle": "Marcus Vance",
    "bank_account": "CA-994-8821-009",
    "target_location": "Pier 4, Warehouse B",
    "stolen_asset": "17th Century Dutch Oil Painting"
}

def create_structure():
    """
    Provisions the physical directory structure mimicking a real investigator's filesystem.

    Returns:
        str: The absolute or relative path to the root of the generated case directory.
    """
    base_dir = f"./{CASE_NAME}"
    folders = ["Financials", "Communications", "Surveillance", "Legal"]

    if not os.path.exists(base_dir):
        os.makedirs(base_dir)

    for folder in folders:
        path = os.path.join(base_dir, folder)
        if not os.path.exists(path):
            os.makedirs(path)

    return base_dir

def generate_case_brief(base_dir):
    """
    Generates the foundational mission document for the QA Tester / Developer.
    This provides the initial context and prompts to feed into the UI to begin the test.
    """
    filepath = os.path.join(base_dir, "00_INVESTIGATION_BRIEF.txt")
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write("=========================================================\n")
        f.write("RESTRICTED POLICE FILE: OPERATION MIDNIGHT\n")
        f.write("=========================================================\n\n")
        f.write("BACKGROUND:\n")
        f.write("A priceless 17th Century Dutch Oil Painting was stolen from the Metro Museum.\n")
        f.write(f"We suspect billionaire {TARGET_ENTITIES['mastermind']} orchestrated the theft.\n\n")
        f.write("OBJECTIVES FOR THE RAG FORENSIC SYSTEM:\n")
        f.write("1. Find out what bank account Victor used to fund this.\n")
        f.write("2. Identify the 'Logistics' coordinator who arranged the transport.\n")
        f.write("3. Identify the 'Muscle' who physically moved the asset.\n")
        f.write("4. Discover the physical location where the painting is currently hidden.\n\n")
        f.write("INSTRUCTIONS:\n")
        f.write("Create Folders in your UI matching this directory structure.\n")
        f.write("Upload the respective files into their folders.\n")
        f.write("Select different folders to see how the AI isolates context!\n")

def generate_financials(base_dir):
    """
    Generates structured Tabular Data (.CSV) to test the memory-safe batching parser.

    TEST INVARIANT: The `AsyncIngestionWorker` uses a sliding semantic window for CSVs.
    By placing the clue exactly at row 250, we verify that the batch-flush mechanism
    (which flushes every 100 rows) correctly embedded the target data into pgvector.
    """
    folder = os.path.join(base_dir, "Financials")
    for i in range(1, 4):
        filepath = os.path.join(folder, f"bank_export_Q{i}.csv")
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(["Transaction_ID", "Date", "Account_From", "Account_To", "Amount", "Notes"])

            for row in range(500):
                # 🚨 THE GOLDEN CLUE (Hidden in the noise of Q2)
                if i == 2 and row == 250:
                    writer.writerow([
                        fake.uuid4(),
                        fake.date_this_year(),
                        TARGET_ENTITIES["bank_account"],
                        "Elena_Rostova_Offshore",
                        "$50,000",
                        f"Wire Transfer - Payment from {TARGET_ENTITIES['mastermind']}"
                    ])
                else:
                    writer.writerow([
                        fake.uuid4(),
                        fake.date_this_year(),
                        fake.iban(),
                        fake.iban(),
                        f"${random.randint(10, 5000)}",
                        fake.word()
                    ])

def generate_communications(base_dir):
    """
    Generates massive unstructured Text (.TXT) to test the `TokenTextSplitter`.

    TEST INVARIANT: 1000 lines of text will vastly exceed the LLM's context window.
    Placing the clue at line 800 forces the vector search to accurately calculate
    Cosine Similarity and pull *only* the specific chunk containing the transaction.
    """
    folder = os.path.join(base_dir, "Communications")
    for i in range(1, 4):
        filepath = os.path.join(folder, f"whatsapp_intercept_device_{i}.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"FORENSIC EXPORT: DEVICE ID #{random.randint(1000, 9999)}\n")
            f.write("=========================================================\n")

            for line in range(1000):
                # 🚨 THE GOLDEN CLUE (Hidden deep in Device 3)
                if i == 3 and line == 800:
                    f.write(f"[2026-06-25 23:14:00] {TARGET_ENTITIES['logistics']}: Listen to me, {TARGET_ENTITIES['muscle']}. Take the {TARGET_ENTITIES['stolen_asset']} and secure it at {TARGET_ENTITIES['target_location']}. Do not fail.\n")
                else:
                    f.write(f"[{fake.date_time_this_year()}] {fake.name()}: {fake.sentence()}\n")

def generate_surveillance(base_dir):
    """
    Generates binary Portable Document Format (.PDF) files to test Apache Tika binary extraction.

    TEST INVARIANT: The LLM must use deductive logic here. The text never explicitly says
    "He stole the painting". It says he carried a crate matching the dimensions.
    """
    folder = os.path.join(base_dir, "Surveillance")
    for i in range(1, 3):
        filepath = os.path.join(folder, f"field_report_unit_{i}.pdf")

        # Instantiate ReportLab Canvas for PDF generation
        c = canvas.Canvas(filepath)
        c.drawString(100, 800, f"OFFICIAL SURVEILLANCE REPORT - UNIT {i}")
        c.drawString(100, 780, f"Date: {fake.date_this_month()}")

        text = c.beginText(100, 730)

        # 🚨 THE GOLDEN CLUE (Deductive Reasoning Check)
        if i == 1:
            text.textLines(f"Subject {TARGET_ENTITIES['muscle']} was observed pacing outside {TARGET_ENTITIES['target_location']}.")
            text.textLines("He was seen carrying a large, rectangular crate matching the dimensions of a painting.")
        else:
            text.textLines(f"Routine observation. {fake.paragraph(nb_sentences=5)}")

        text.textLines(f"Additional notes: {fake.paragraph(nb_sentences=10)}")
        c.drawText(text)
        c.save()

def generate_legal(base_dir):
    """
    Generates Microsoft Word (.DOCX) files to test structured XML extraction.

    TEST INVARIANT: This provides the final connective tissue for the investigation,
    linking the specific bank account back to the mastermind.
    """
    folder = os.path.join(base_dir, "Legal")
    filepath = os.path.join(folder, "search_warrant_001.docx")

    doc = Document()
    doc.add_heading('Search Warrant Authorized', 0)

    # 🚨 THE GOLDEN CLUE (Entity Resolution Check)
    doc.add_paragraph(f"By order of the court, accounts linked to {TARGET_ENTITIES['bank_account']} are frozen.")
    doc.add_paragraph(f"These accounts are confirmed to be the primary slush fund of {TARGET_ENTITIES['mastermind']}.")

    # Pad with irrelevant text to test chunking
    doc.add_paragraph(fake.text(max_nb_chars=1000))
    doc.save(filepath)

if __name__ == "__main__":
    print("Initializing Forensic Data Generator...")

    # 1. Bootstraps the folder topology
    base = create_structure()

    # 2. Generates the synthesized artifacts
    generate_case_brief(base)
    generate_financials(base)
    generate_communications(base)
    generate_surveillance(base)
    generate_legal(base)

    print(f"\nSUCCESS! Generated highly structured investigation case in: ./{CASE_NAME}")
    print("Folders created: Financials, Communications, Surveillance, Legal")
    print("Read '00_INVESTIGATION_BRIEF.txt' for your instructions.")